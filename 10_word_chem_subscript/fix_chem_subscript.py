"""
fix_chem_subscript.py
Word docx 내 화학식 subscript/superscript 자동 수정 스크립트

지원 패턴:
  Subscript:  MnO2, PtO2, CeO2, TiO2, PtF6, NH3, O2, N2, CO2, H2O, SO4, ...
  Superscript: Pt6+, Fe3+, Co2+, ... (원소 뒤 숫자+charge)

사용법:
  python fix_chem_subscript.py <input.docx> [output.docx]
  output 생략 시 원본 덮어씀 (백업 자동 생성)
"""

import re
import sys
import shutil
from docx import Document
from lxml import etree
from copy import deepcopy

# XML namespace constants
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

# ──────────────────────────────────────────────
# Pattern definitions
# (fullmatch_regex, split_function)
# split_function: matched_text -> [(text, fmt), ...]
#   fmt: None=normal, 'subscript', 'superscript'
# ──────────────────────────────────────────────

PATTERN_DEFS = [
    # Superscript: 원소 + 숫자 + charge (Pt6+, Fe3+, Co2+, Mn4+, ...)
    (r'[A-Z][a-z]?\d+[+\-]', lambda m: [(re.match(r'[A-Z][a-z]?', m).group(), None),
                                          (m[re.match(r'[A-Z][a-z]?', m).end():], 'superscript')]),
    # Subscript: 화학식 내 숫자 (MnO2, PtO2, CeO2, TiO2, H2O, CO2, ...)
    # 원소2글자 + 원소1글자 + 숫자
    (r'[A-Z][a-z]?[A-Z]?[a-z]?\d+', lambda m: [(re.match(r'[A-Z][a-z]?[A-Z]?[a-z]?', m).group(), None),
                                                  (m[re.match(r'[A-Z][a-z]?[A-Z]?[a-z]?', m).end():], 'subscript')]),
    # 단독 O2, N2, H2 등 (앞에 영문자 없을 때)
    (r'(?<![A-Za-z])[ONHC]\d+(?!\d)', lambda m: [(m[0], None), (m[1:], 'subscript')]),
]

# Combined regex for finditer (order matters: specific first)
COMBINED_RE = re.compile(
    r'[A-Z][a-z]?\d+[+\-]'           # superscript: Pt6+, Fe3+
    r'|MnO\d+|PtO\d+|CeO\d+|TiO\d+'  # known oxide formulas
    r'|PtF\d+|CoF\d+|FeF\d+'          # fluorides
    r'|NH\d+|CH\d+|SH\d+'             # hydrides
    r'|SO\d+|NO\d+|CO\d+'             # oxyanions/molecules
    r'|H2O|H2O2'                       # water, peroxide
    r'|(?<![A-Za-z])[ONHC]\d+(?!\d)'  # standalone O2, N2, H2, C2
)


def match_to_segments(matched_text):
    """Map a matched chemical string to [(text, fmt), ...] segments."""
    for pattern, split_func in PATTERN_DEFS:
        if re.fullmatch(pattern, matched_text):
            return split_func(matched_text)
    return [(matched_text, None)]


def get_segments(text):
    """Split text into segments, tagging chemical numbers for sub/superscript."""
    segments = []
    last_end = 0
    for match in COMBINED_RE.finditer(text):
        if match.start() > last_end:
            segments.append((text[last_end:match.start()], None))
        segments.extend(match_to_segments(match.group()))
        last_end = match.end()
    if last_end < len(text):
        segments.append((text[last_end:], None))
    return segments


def process_run(run_elem):
    """Process a single <w:r> element. Returns number of formatting changes."""
    t_elem = run_elem.find(f'{W}t')
    if t_elem is None or not t_elem.text:
        return 0

    segments = get_segments(t_elem.text)
    if not any(fmt for _, fmt in segments):
        return 0

    parent = run_elem.getparent()
    idx = list(parent).index(run_elem)
    changes = 0

    new_elements = []
    for seg_text, fmt in segments:
        if not seg_text:
            continue
        new_run = deepcopy(run_elem)
        new_t = new_run.find(f'{W}t')
        new_t.text = seg_text
        new_t.set(XML_SPACE, 'preserve')

        if fmt:
            rpr = new_run.find(f'{W}rPr')
            if rpr is None:
                rpr = etree.Element(f'{W}rPr')
                new_run.insert(0, rpr)
            for va in rpr.findall(f'{W}vertAlign'):
                rpr.remove(va)
            va = etree.SubElement(rpr, f'{W}vertAlign')
            va.set(f'{W}val', fmt)
            changes += 1
        else:
            rpr = new_run.find(f'{W}rPr')
            if rpr is not None:
                for va in rpr.findall(f'{W}vertAlign'):
                    rpr.remove(va)

        new_elements.append(new_run)

    for i, el in enumerate(new_elements):
        parent.insert(idx + i, el)
    parent.remove(run_elem)
    return changes


def collect_runs(elem):
    """Collect <w:r> elements from paragraph XML, including inside hyperlinks."""
    runs = []
    for child in list(elem):
        tag = etree.QName(child.tag).localname if '}' in child.tag else child.tag
        if tag == 'r':
            runs.append(child)
        elif tag in ('hyperlink', 'ins', 'smartTag', 'sdtContent', 'fldSimple'):
            runs.extend(collect_runs(child))
    return runs


def fix_document(input_path, output_path=None):
    """Main function: fix subscript/superscript in a docx file."""
    if output_path is None:
        output_path = input_path

    # Backup
    backup_path = input_path.replace('.docx', '_backup.docx')
    shutil.copy2(input_path, backup_path)
    print(f"Backup: {backup_path}")

    doc = Document(input_path)
    total = 0

    # Body paragraphs
    for para in doc.paragraphs:
        for run_elem in collect_runs(para._element):
            total += process_run(run_elem)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run_elem in collect_runs(para._element):
                        total += process_run(run_elem)

    doc.save(output_path)
    print(f"Changes: {total}")
    print(f"Saved: {output_path}")
    return total


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python fix_chem_subscript.py <input.docx> [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    fix_document(input_file, output_file)
